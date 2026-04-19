#!/usr/bin/env python3
import argparse, csv, json, logging, os, re, signal, sys, time, warnings, contextlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from urllib.parse import unquote

import pandas as pd

warnings.filterwarnings("ignore")
logging.disable(logging.CRITICAL)

STRUCT = {".csv", ".json", ".parquet"}
DOCS   = {".pdf", ".doc", ".docx", ".rtf", ".xls", ".xlsx"}
WEB    = {".html", ".htm"}
IMAGES = {".tif", ".tiff", ".jpeg", ".jpg", ".png", ".gif"}
VIDEO  = {".mp4"}
PLAIN  = {".txt", ".md", ".markdown"}
ALL    = STRUCT | DOCS | WEB | IMAGES | VIDEO | PLAIN

ENCODINGS      = ("utf-8", "utf-8-sig", "cp1251", "latin-1")
MONTHS         = ("jan","feb","mar","apr","may","jun","jul","aug","sep","oct","nov","dec")
MAX_TEXT_BYTES = 800_000
MAX_CSV_ROWS   = 10_000
MAX_PDF_PAGES  = 15
MAX_HTML_BYTES = 150_000

RE_EMAIL    = re.compile(r"\b[a-zA-Z0-9][a-zA-Z0-9_.+-]*@[a-zA-Z0-9-]{2,}\.[a-zA-Z]{2,}\b", re.I)
RE_PHONE    = re.compile(r"(?<!\d)(?:\+7|8)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}(?!\d)")
RE_PASSPORT = re.compile(r"(?<!\d)\d{4}\s?\d{6}(?!\d)")
RE_SNILS    = re.compile(r"\b\d{3}[\s\-]\d{3}[\s\-]\d{3}[\s\-]\d{2}\b")
RE_SNILS_CTX = re.compile(r"(?i)(?:снилс|страховой|пенсионн)[^\d]{0,30}(\d{3}[\s\-]?\d{3}[\s\-]?\d{3}[\s\-]?\d{2}|\d{11})")
RE_IDENTITY  = re.compile(r"(?i)(?:фамилия|имя|отчество|фио|дата.рожден|место.рожден)", re.I)
RE_INN12    = re.compile(r"(?<!\d)\d{12}(?!\d)")
RE_INN10_C  = re.compile(r"(?i)(?:инн|inn)[:\s№#]*(\d{10})\b")
RE_CARD     = re.compile(r"\b(?:\d{4}[\s\-]){3}\d{4}\b|\b(?:\d{4}[\s\-]){2}\d{6,8}\b|\b\d{16}\b")
RE_DOB      = re.compile(r"\b(?:0[1-9]|[12]\d|3[01])[./-](?:0[1-9]|1[0-2])[./-](?:19|20)\d{2}\b")
RE_BIK      = re.compile(r"\b04\d{7}\b")
RE_ACCOUNT  = re.compile(r"\b\d{20}\b")

PD_COL = re.compile(
    r"(?i)^(?:"
    r"(?:first|last|full|middle)[_\s]?name|name|имя|фамилия|отчество|фио|"
    r"phone|телефон|mobile|"
    r"email|e[_\-]?mail|почта|"
    r"address|адрес|"
    r"birth|рожден|dob|"
    r"passport|паспорт|"
    r"snils|снилс|"
    r"inn|инн|"
    r"gender|пол|"
    r"customer|client|subscriber|employee|сотрудник|клиент|пользователь|user"
    r")$", re.I | re.UNICODE
)

PD_FNAME = re.compile(
    r"(?i)(?:"
    r"anketa|анкета|"
    r"zayavka|zayavleni|заявк|заявлени|"
    r"soglasie|согласи|"
    r"doverennost|доверенност|"
    r"raspiska|расписк|"
    r"perepiska|переписк|"
    r"dogovor|договор|"
    r"customers?|clients?|"
    r"subscribers?|"
    r"physical|phys_|"
    r"employee|personnel|"
    r"patient|"
    r"logistics|"
    r"sales|"
    r"stores?|"
    r"incidents?|"
    r"sample|"
    r"pii|pd_scan|personaldaten|"
    r"vypiska|выписк|"
    r"reestr|реестр|"
    r"company|"
    r"contacts?|kontakt"
    r")", re.I
)

PDF_PUBLIC = re.compile(
    r"(?i)(?:"
    r"pravil|"
    r"prot(?:okol)?\d*[_.\s]|"
    r"otchet|отчёт|отчет|"
    r"global|portrait|festival|"
    r"programm[au]|programma|program.?mag|"
    r"gender.equal|"
    r"fizich|physical.?cultur|"
    r"samoobsled|самообслед|"
    r"priema|приём|pravila.?prie|"
    r"bsm|pp.?bsm|"
    r"indeksaci|индексац|"
    r"rasporj|расп\d|rasporiageni|"
    r"version.?\d{6}|"
    r"outlook\d|"
    r"progress.?report|"
    r"self.?portrait|"
    r"cvdrkased|"
    r"syllabus|course.?sample|"
    r"fin.?result|finansov|"
    r"scientific.?research|"
    r"urban.?design|urban.?develop|"
    r"environment.?problem|"
    r"ecology.?of|"
    r"territory.?brand|"
    r"key.?concept"
    r")", re.I
)

SPECIAL: dict[str, list[str]] = {
    "здоровье":       ["история болезни", "амбулаторная карта", "диагноз", "инвалидность"],
    "биометрия":      ["отпечатки пальцев", "радужная оболочка", "биометрические данные"],
    "религия":        ["вероисповедание", "религиозные убеждения"],
    "политика":       ["политические убеждения", "партийная принадлежность"],
    "национальность": ["расовая принадлежность", "национальная принадлежность", "этническая принадлежность"],
    "судимость":      ["судимость", "уголовное преследование"],
}

_W10  = [2, 4, 10, 3, 5, 9, 4, 6, 8]
_W12A = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
_W12B = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8]


def luhn(s: str) -> bool:
    d = [int(c) for c in re.sub(r"\D", "", s)]
    if not 13 <= len(d) <= 19:
        return False
    return (sum(d[::-2]) + sum(sum(divmod(x * 2, 10)) for x in d[-2::-2])) % 10 == 0


def snils_ok(s: str) -> bool:
    d = re.sub(r"\D", "", s)
    if len(d) != 11:
        return False
    c = sum(int(d[i]) * (9 - i) for i in range(9))
    if c > 101: c %= 101
    if c == 100: c = 0
    return c == int(d[9:])


def inn_ok(s: str) -> bool:
    d = [int(c) for c in re.sub(r"\D", "", s)]
    if len(d) == 12:
        return (sum(d[i] * _W12A[i] for i in range(10)) % 11 % 10 == d[10] and
                sum(d[i] * _W12B[i] for i in range(11)) % 11 % 10 == d[11])
    if len(d) == 10:
        return sum(d[i] * _W10[i] for i in range(9)) % 11 % 10 == d[9]
    return False


def _read(path: str, limit: int = MAX_TEXT_BYTES) -> str:
    for enc in ENCODINGS:
        try:
            with open(path, encoding=enc, errors="strict") as f:
                return f.read(limit)
        except Exception:
            continue
    try:
        with open(path, "rb") as f:
            return f.read(limit).decode("utf-8", errors="replace")
    except Exception:
        return ""


def _df_text(df: pd.DataFrame) -> str:
    df = df.fillna("").astype(str)
    return (" | ".join(str(c) for c in df.columns) + "\n" +
            df.to_string(index=False, header=False))[:MAX_TEXT_BYTES]


def _flatten(obj, depth=0) -> list:
    if depth > 5:
        return []
    if isinstance(obj, dict):
        out = []
        for k, v in obj.items():
            out.append(str(k))
            out.extend(_flatten(v, depth + 1))
        return out
    if isinstance(obj, list):
        out = []
        for v in obj[:2000]:
            out.extend(_flatten(v, depth + 1))
        return out
    return [str(obj)]


def _quick_reject(path: str) -> bool:
    fname = unquote(Path(path).name).lower()
    if PDF_PUBLIC.search(fname):
        return True
    sz = Path(path).stat().st_size
    if Path(path).suffix.lower() in IMAGES and sz < 5_000:
        return True
    if Path(path).suffix.lower() == ".pdf" and sz > 25_000_000:
        return True
    return False


def _read_csv_fast(path: str) -> tuple[str, list[str]]:
    for enc in ENCODINGS:
        for sep in (",", ";", "\t", "|"):
            try:
                df_head = pd.read_csv(path, nrows=1, encoding=enc, sep=sep, on_bad_lines="skip")
                if len(df_head.columns) > 1:
                    cols = list(df_head.columns.astype(str))
                    df = pd.read_csv(path, nrows=MAX_CSV_ROWS, encoding=enc, sep=sep,
                                     on_bad_lines="skip", low_memory=False)
                    return _df_text(df), cols
            except Exception:
                continue
    try:
        df = pd.read_csv(path, nrows=MAX_CSV_ROWS, sep=None, engine="python", on_bad_lines="skip")
        if len(df.columns) > 1:
            return _df_text(df), list(df.columns.astype(str))
    except Exception:
        pass
    return "", []


def _read_pdf_fast(path: str) -> str:
    parts: list[str] = []
    total_len = 0
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:MAX_PDF_PAGES]:
                try:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
                        total_len += len(t)
                        if total_len > MAX_TEXT_BYTES:
                            break
                except Exception:
                    continue
    except Exception:
        pass
    if not parts:
        try:
            from pypdf import PdfReader
            reader = PdfReader(path, strict=False)
            for page in reader.pages[:MAX_PDF_PAGES]:
                try:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
                        total_len += len(t)
                        if total_len > MAX_TEXT_BYTES:
                            break
                except Exception:
                    continue
        except Exception:
            pass
    return "\n".join(parts)[:MAX_TEXT_BYTES]


def _ocr_image(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image, ImageOps
    except Exception:
        return ""
    try:
        with Image.open(path) as img:
            gray = img.convert("L")
            w, h = gray.size
            if max(w, h) < 1200:
                scale = max(1, 1400 // max(w, h))
                gray = gray.resize((w * scale, h * scale))
            enhanced = ImageOps.autocontrast(gray)
            for cfg in ("--oem 1 --psm 6", "--oem 1 --psm 1", "--oem 1 --psm 3"):
                try:
                    text = pytesseract.image_to_string(enhanced, lang="rus+eng", config=cfg, timeout=60)
                    if text and len(text.strip()) > 20:
                        return text
                except (RuntimeError, Exception):
                    continue
            try:
                return pytesseract.image_to_string(gray, lang="rus+eng", timeout=30)
            except Exception:
                return ""
    except Exception:
        return ""


def get_text(path: str, skip_ocr: bool) -> tuple[str, list[str]]:
    ext  = Path(path).suffix.lower()
    cols: list[str] = []
    text = ""
    try:
        if ext == ".csv":
            text, cols = _read_csv_fast(path)

        elif ext == ".parquet":
            df   = pd.read_parquet(path).head(MAX_CSV_ROWS)
            cols = list(df.columns.astype(str))
            text = _df_text(df)

        elif ext in (".xls", ".xlsx"):
            for engine in ("openpyxl", "xlrd"):
                try:
                    xls  = pd.ExcelFile(path, engine=engine)
                    df   = xls.parse(sheet_name=xls.sheet_names[0], nrows=MAX_CSV_ROWS)
                    cols = list(df.columns.astype(str))
                    text = _df_text(df)
                    break
                except Exception:
                    continue

        elif ext == ".json":
            try:
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, list) and data and isinstance(data[0], dict):
                    cols = list(data[0].keys())
                text = "\n".join(_flatten(data))[:MAX_TEXT_BYTES]
            except Exception:
                text = _read(path)

        elif ext == ".pdf":
            text = _read_pdf_fast(path)

        elif ext == ".docx":
            try:
                from docx import Document
                doc   = Document(path)
                parts = [p.text for p in doc.paragraphs if p.text.strip()][:100]
                for tbl in doc.tables[:10]:
                    for row in tbl.rows[:20]:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            parts.append(" | ".join(cells))
                text = "\n".join(parts)[:MAX_TEXT_BYTES]
            except Exception:
                pass

        elif ext == ".rtf":
            try:
                from striprtf.striprtf import rtf_to_text
                with open(path, encoding="utf-8", errors="ignore") as f:
                    text = rtf_to_text(f.read(MAX_TEXT_BYTES))[:MAX_TEXT_BYTES]
            except Exception:
                pass

        elif ext in WEB:
            try:
                from bs4 import BeautifulSoup
                raw  = _read(path, MAX_HTML_BYTES)
                soup = BeautifulSoup(raw, "html.parser")
                for tag in soup(["script", "style", "noscript", "nav", "footer", "header", "aside"]):
                    tag.decompose()
                text = soup.get_text(separator=" ", strip=True)[:MAX_TEXT_BYTES]
            except Exception:
                text = _read(path, MAX_HTML_BYTES)

        elif ext in PLAIN:
            text = _read(path, MAX_TEXT_BYTES)

        elif ext in IMAGES and not skip_ocr:
            text = _ocr_image(path)

        elif ext in VIDEO and not skip_ocr:
            try:
                import cv2, pytesseract
                from PIL import Image, ImageOps
                cap   = cv2.VideoCapture(path)
                total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
                idxs  = [0, total // 2, max(0, total - 1)] if total > 2 else [0]
                parts = []
                for idx in set(idxs):
                    cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
                    ok, frame = cap.read()
                    if not ok:
                        continue
                    img  = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                    gray = ImageOps.autocontrast(img.convert("L"))
                    try:
                        t = pytesseract.image_to_string(gray, lang="rus+eng",
                                                        config="--oem 1 --psm 1", timeout=15)
                    except Exception:
                        t = ""
                    if t.strip():
                        parts.append(t)
                cap.release()
                text = "\n".join(parts)[:MAX_TEXT_BYTES]
            except Exception:
                pass

    except Exception:
        pass

    return text, cols


def find_pd(text: str) -> dict:
    r: dict[str, int] = {}
    if not text:
        return r

    snils = [m for m in RE_SNILS.findall(text) if snils_ok(m)]
    if not snils:
        snils = [m for m in RE_SNILS_CTX.findall(text) if snils_ok(re.sub(r"\D", "", m)[:11] if len(re.sub(r"\D", "", m)) == 11 else m)]
    if snils:
        r["снилс"] = len(snils)

    inn12 = [m for m in RE_INN12.findall(text) if inn_ok(m)]
    if inn12:
        r["инн_физ"] = len(inn12)

    cards = [m for m in RE_CARD.findall(text) if luhn(m)]
    if cards:
        r["карта"] = len(cards)

    inn10 = [m for m in RE_INN10_C.findall(text) if inn_ok(m)]
    if inn10:
        r["инн_юр"] = len(inn10)

    if RE_BIK.search(text) and RE_ACCOUNT.search(text):
        r["счёт"] = len(RE_ACCOUNT.findall(text))

    emails = RE_EMAIL.findall(text)
    if emails:
        r["email"] = len(emails)

    phones = RE_PHONE.findall(text)
    if phones:
        r["телефон"] = len(phones)

    dobs = RE_DOB.findall(text)
    if dobs:
        r["дата_рождения"] = len(dobs)

    lo = text.lower()
    if any(w in lo for w in ("паспорт", "passport", "серия", "series")):
        pp = RE_PASSPORT.findall(text)
        if pp:
            r["паспорт"] = len(pp)

    for cat, kws in SPECIAL.items():
        if any(kw in lo for kw in kws):
            r[f"спец:{cat}"] = 1

    if RE_IDENTITY.search(text):
        r["identity_context"] = 1

    return r


def cols_have_pd(cols: list[str]) -> list[str]:
    return [c for c in cols if PD_COL.match(c.strip())]


def is_pd(path: str, text: str, cols: list[str], found: dict) -> tuple[bool, str]:
    ext        = Path(path).suffix.lower()
    fname      = unquote(Path(path).name)
    fname_lo   = fname.lower()
    fname_hint = bool(PD_FNAME.search(fname))

    strong   = found.get("снилс", 0) + found.get("инн_физ", 0) + found.get("карта", 0)
    specials = [k for k in found if k.startswith("спец:")]
    email    = found.get("email", 0)
    phone    = found.get("телефон", 0)
    dob      = found.get("дата_рождения", 0)
    passport = found.get("паспорт", 0)
    contacts = int(email > 0) + int(phone > 0)

    if ext in {".pdf", ".docx", ".doc", ".rtf"}:
        if PDF_PUBLIC.search(fname):
            return False, ""
        if any(w in fname_lo for w in ("politika", "policy", "privacy",
                                       "rasporiageni", "распоряжени",
                                       "prilozhenie", "приложение",
                                       "условия использования",
                                       "конкурс", "konkurs")):
            return False, ""
        stem = Path(path).stem
        if re.fullmatch(r"[А-ЯЁA-Z]\d+", stem, re.I):
            return False, ""
        if re.fullmatch(r"\d{4}[-–]\d{4}[-–]\d{2}[-–]\d{2}", stem):
            return False, ""
        if re.fullmatch(r"\d{2,4}[A-Z]{2,}\d+", stem, re.I):
            return False, ""

    if strong > 0:
        return True, f"сильный_ид:{','.join(k for k in ('снилс','инн_физ','карта') if found.get(k,0))}"

    if specials and (contacts > 0 or dob > 0 or passport > 0):
        return True, f"спец_кат:{'+'.join(s.split(':')[1] for s in specials)}"

    if ext in STRUCT | {".xls", ".xlsx"}:
        pd_cols = cols_have_pd(cols)
        if pd_cols:
            return True, f"pd_cols:{','.join(pd_cols[:4])}"
        if fname_hint and (contacts > 0 or dob > 0 or len(cols) > 2):
            return True, "struct_name"
        if ext == ".json":
            if contacts >= 1 and dob:
                return True, "json:contact+dob"
            if contacts >= 1 and passport:
                return True, "json:contact+passport"
            if fname_hint and contacts >= 1:
                return True, "json_name+contact"
        return False, ""

    if ext in PLAIN:
        if fname_hint and (contacts > 0 or dob > 0 or passport > 0):
            return True, "txt_name+pd"
        if passport and contacts > 0:
            return True, "txt:passport+contact"
        if contacts == 2 and dob:
            return True, "txt:email+phone+dob"
        if strong > 0:
            return True, "md:strong_id"
        return False, ""

    if ext in WEB:
        if strong > 0 and contacts > 0:
            return True, "html:strong_id+contact"
        if passport and contacts > 0 and dob:
            return True, "html:pass+contact+dob"
        if contacts == 2 and dob:
            return True, "html:2contacts+dob"
        return False, ""

    if ext in {".pdf", ".docx", ".doc", ".rtf"}:
        if ext == ".rtf" and fname_hint:
            return True, "rtf_name"
        if fname_hint and (contacts > 0 or passport > 0 or dob > 0):
            return True, "doc_name+pd"
        if passport and contacts > 0 and dob:
            return True, "doc:pass+contact+dob"
        if passport and contacts == 2:
            return True, "doc:pass+2contacts"
        return False, ""

    if ext in IMAGES | VIDEO:
        if strong > 0:
            return True, "ocr:strong_id"
        identity = found.get("identity_context", 0)
        if passport and (contacts > 0 or dob > 0 or identity > 0):
            return True, "ocr:pass+context"
        if dob > 0 and identity > 0 and (contacts > 0 or passport > 0 or strong > 0):
            return True, "ocr:dob+identity"
        if contacts == 2 and (dob > 0 or identity > 0 or passport > 0):
            return True, "ocr:contacts+pd"
        if specials and (dob > 0 or identity > 0):
            return True, "ocr:specials"
        return False, ""

    return False, ""


def classify_uz(found: dict) -> str:
    specials = [k for k in found if k.startswith("спец:")]
    strong   = found.get("снилс", 0) + found.get("инн_физ", 0) + found.get("карта", 0)
    payment  = found.get("карта", 0) + found.get("счёт", 0)
    contacts = found.get("email", 0) + found.get("телефон", 0)
    if specials:
        return "УЗ-1"
    if payment >= 10 or strong >= 20:
        return "УЗ-2"
    if strong > 0 or contacts >= 50:
        return "УЗ-3"
    return "УЗ-4"


def process_file(args_tuple: tuple) -> dict | None:
    path, skip_ocr = args_tuple
    try:
        if _quick_reject(path):
            return None
        text, cols = get_text(path, skip_ocr)
        found      = find_pd(text)
        ok, reason = is_pd(path, text, cols, found)
        if not ok:
            return None
        stat     = Path(path).stat()
        dt       = datetime.fromtimestamp(stat.st_mtime)
        time_str = f"{MONTHS[dt.month - 1].title()} {dt.day:02d} {dt.strftime('%H:%M')}"
        cats     = ", ".join(k for k in found if found[k] > 0)
        return {
            "name":   unquote(Path(path).name),
            "size":   stat.st_size,
            "time":   time_str,
            "path":   path,
            "uz":     classify_uz(found),
            "cats":   cats,
            "reason": reason,
        }
    except Exception:
        return None


def find_files(root: Path, skip_ocr: bool):
    struct, heavy = [], []
    for p in root.rglob("*"):
        if not p.is_file() or p.name.startswith("."):
            continue
        ext = p.suffix.lower()
        if ".ipynb_checkpoints" in p.parts and ext == ".ipynb":
            continue
        if ext not in ALL:
            continue
        if skip_ocr and ext in IMAGES | VIDEO:
            continue
        (struct if ext in STRUCT else heavy).append(str(p))
    return struct, heavy


def write_result_csv(rows: list) -> Path:
    out = Path.cwd() / "result.csv"
    with contextlib.suppress(OSError):
        out.unlink()
    with open(out, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["size", "time", "name"], lineterminator="\n")
        w.writeheader()
        for r in rows:
            w.writerow({"size": r["size"], "time": r["time"], "name": r["name"]})
    return out


def write_report(rows: list, out: Path) -> None:
    with contextlib.suppress(OSError):
        out.unlink()
    with open(out, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["name", "uz", "cats", "reason", "size", "time", "path"],
                           delimiter=";", lineterminator="\n")
        w.writeheader()
        for r in rows:
            w.writerow(r)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-dir",  default=os.environ.get("INPUT_DIR", "/DATA"))
    ap.add_argument("--output",     default="report_pd_scan.csv")
    ap.add_argument("--max-files",  type=int, default=0)
    ap.add_argument("--threads",    type=int, default=min(16, os.cpu_count() or 4))
    ap.add_argument("--skip-ocr",   action="store_true")
    args = ap.parse_args()

    root = Path(args.input_dir).expanduser()
    if not root.exists():
        print(f"[!] Папка не найдена: {root}", flush=True)
        sys.exit(1)

    print(f"[СТАРТ] {root}", flush=True)
    struct, heavy = find_files(root, args.skip_ocr)
    all_f = struct + heavy
    if args.max_files > 0:
        all_f  = all_f[:args.max_files]
        s      = set(all_f)
        struct = [p for p in struct if p in s]
        heavy  = [p for p in heavy  if p in s]
    total = len(all_f)

    if not total:
        print("[!] Файлов не найдено, проверьте --input-dir", flush=True)
        sys.exit(1)

    ocr_info = "с OCR" if not args.skip_ocr else "без OCR"
    print(f"[INFO] Файлов: {total} | структур: {len(struct)} | тяжёлых: {len(heavy)} | {ocr_info}", flush=True)

    results: list[dict] = []
    done = 0
    t0   = time.time()

    def _batch(files: list, workers: int, timeout: int = 120) -> None:
        nonlocal done
        with ThreadPoolExecutor(max_workers=max(1, workers)) as ex:
            futs = {ex.submit(process_file, (p, args.skip_ocr)): p for p in files}
            for fut in as_completed(futs):
                done += 1
                try:
                    row = fut.result(timeout=timeout)
                    if row:
                        results.append(row)
                except Exception:
                    pass
                if done % 50 == 0 or done == total:
                    el  = time.time() - t0
                    spd = done / el if el > 0 else 0
                    eta = int((total - done) / spd) if spd > 0 else 0
                    print(
                        f"  [{done:>4}/{total}] ПДн: {len(results):>3} | {spd:.1f} ф/с | ≈{eta}с",
                        flush=True
                    )

    fast = struct + [p for p in heavy if Path(p).suffix.lower() in PLAIN | WEB]
    docs = [p for p in heavy if Path(p).suffix.lower() in {".pdf", ".docx", ".rtf", ".doc", ".xls", ".xlsx"}]
    ocr  = [p for p in heavy if Path(p).suffix.lower() in IMAGES | VIDEO]

    _batch(fast, workers=args.threads,                                      timeout=30)
    _batch(docs, workers=max(1, args.threads // 2),                         timeout=60)
    _batch(ocr,  workers=max(1, min(args.threads, os.cpu_count() or 4)),    timeout=180)

    elapsed = time.time() - t0
    print(f"\n[ИТОГ] Обработано: {total} | ПДн: {len(results)} | {elapsed:.1f}с", flush=True)

    if not results:
        print("[РЕЗУЛЬТАТ] Файлов с ПД не найдено", flush=True)
        with contextlib.suppress(OSError):
            (Path.cwd() / "result.csv").unlink()
        sys.exit(0)

    results.sort(key=lambda r: r["name"])
    rc = write_result_csv(results)
    write_report(results, Path(args.output).expanduser().resolve())

    uz: dict[str, int] = {}
    for r in results:
        uz[r["uz"]] = uz.get(r["uz"], 0) + 1
    print(f"[УЗ]  {uz}", flush=True)
    print(f"[result.csv] → {rc}", flush=True)


if __name__ == "__main__":
    signal.signal(signal.SIGINT,  lambda *_: sys.exit(0))
    signal.signal(signal.SIGTERM, lambda *_: sys.exit(0))
    main()
